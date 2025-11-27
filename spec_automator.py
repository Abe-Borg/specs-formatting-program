import re
import os
import json
from docx import Document

# ==========================================
# 🛠️ SETUP & CONFIGURATION
# ==========================================
BASE_DIR = os.getcwd()
CONFIG_FILE = os.path.join(BASE_DIR, "project_config.json")

# Default Fallback Settings (Used if no JSON is found)
DEFAULT_CONFIG = {
    "styles": {
        "Title": "Heading 1", 
        "Part": "Heading 2", 
        "Article": "Heading 3"
    },
    "list_levels": ["List Bullet", "List Bullet 2", "List Bullet 3"],
    "options": {
        "strip_heading_numbers": True, 
        "strip_list_labels": True
    }
}

def load_config():
    if os.path.exists(CONFIG_FILE):
        print(f"⚙️  Loading configuration from {CONFIG_FILE}...")
        try:
            with open(CONFIG_FILE, 'r') as f:
                return json.load(f)
        except Exception as e:
            print(f"⚠️  Error reading config file: {e}. Using defaults.")
    else:
        print("⚠️  No 'project_config.json' found. Using default Standard Word styles.")
    return DEFAULT_CONFIG

# Load Config Immediately
CONFIG = load_config()
STYLE_MAP = CONFIG["styles"]
LIST_MAP = CONFIG["list_levels"]
OPTIONS = CONFIG["options"]

# ==========================================
# STEP 1: EXTRACT TO MARKDOWN
# ==========================================
def extract_master_to_markdown(docx_path, md_path):
    """
    Converts Office Master -> Clean Markdown.
    This neutralizes hidden formatting and converts auto-numbers to text.
    """
    if not os.path.exists(docx_path):
        print(f"❌ Error: Master file {docx_path} not found.")
        return

    doc = Document(docx_path)
    md_lines = []

    print(f"1️⃣  Extracting content from {docx_path}...")

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text: continue

        # Regex detection for CSI Structure
        if text.upper().startswith("SECTION"):
            md_lines.append(f"# {text}")
        elif text.upper().startswith("PART") and "-" in text:
            md_lines.append(f"\n## {text}")
        elif re.match(r'^\d+\.\d+\s', text):
            md_lines.append(f"\n### {text}")
        else:
            # Indentation Detection
            if re.match(r'^[A-Z]\.\s', text):       md_lines.append(f"- {text}")
            elif re.match(r'^\d+\.\s', text):       md_lines.append(f"  - {text}")
            elif re.match(r'^[a-z]\.\s', text):     md_lines.append(f"    - {text}")
            elif re.match(r'^\d+\)\s', text):       md_lines.append(f"      - {text}")
            elif re.match(r'^[a-z]\)\s', text):     md_lines.append(f"        - {text}")
            else:                                   md_lines.append(f"  {text}")

    with open(md_path, "w", encoding="utf-8") as f:
        f.write("\n".join(md_lines))
    print(f"   Converted to Markdown: {md_path}")

# ==========================================
# STEP 2: REBUILD DOCX
# ==========================================
def rebuild_from_markdown(md_path, template_path, output_path):
    """
    Reads Markdown -> Pours into Template Copy -> Saves Final Docx
    """
    if not os.path.exists(template_path):
        print(f"❌ Error: Template {template_path} not found.")
        return

    print(f"2️⃣  Applying styles from {template_path}...")
    
    doc = Document(template_path)
    
    # Clear Body (Keep Headers/Footers/Margins)
    for element in list(doc.element.body):
        if element.tag.endswith('sectPr'): continue
        doc.element.body.remove(element)

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    # --- HELPERS ---
    def add_safe_paragraph(text, style_name):
        try:
            doc.add_paragraph(text, style=style_name)
        except KeyError:
            print(f"⚠️  Style '{style_name}' not found. Using 'Normal'.")
            doc.add_paragraph(text, style='Normal')

    def clean_heading_label(text):
        if OPTIONS.get("strip_heading_numbers"):
            return re.sub(r'^\d+\.\d+\s+', '', text)
        return text

    def clean_list_label(text):
        if OPTIONS.get("strip_list_labels"):
            return re.sub(r'^[A-Za-z0-9]+[\.\)]\s+', '', text)
        return text

    # --- PROCESSING LOOP ---
    previous_line_was_title = False

    for line in lines:
        raw_text = line.rstrip()
        stripped_text = raw_text.strip()
        
        if not stripped_text: continue

        # SPECIAL: End of Section -> Force Title Style
        if stripped_text.upper().startswith("END OF SECTION"):
            add_safe_paragraph(stripped_text, STYLE_MAP["Title"])
            previous_line_was_title = False
            continue

        # HEADERS
        if stripped_text.startswith("# "):
            clean_text = stripped_text.replace("# ", "")
            add_safe_paragraph(clean_text, STYLE_MAP["Title"])
            previous_line_was_title = True 
            
        elif stripped_text.startswith("## "):
            clean_text = stripped_text.replace("## ", "")
            add_safe_paragraph(clean_text, STYLE_MAP["Part"])
            previous_line_was_title = False
            
        elif stripped_text.startswith("### "):
            text_no_hash = stripped_text.replace("### ", "")
            clean_text = clean_heading_label(text_no_hash)
            add_safe_paragraph(clean_text, STYLE_MAP["Article"])
            previous_line_was_title = False

        # LISTS
        elif stripped_text.startswith("- "):
            dash_index = raw_text.find("-")
            indent_level = dash_index // 2 
            
            text_without_dash = stripped_text[2:]
            clean_content = clean_list_label(text_without_dash)

            # Pick style based on indent depth
            if indent_level < len(LIST_MAP):
                target_style = LIST_MAP[indent_level]
            else:
                target_style = LIST_MAP[-1]
            
            add_safe_paragraph(clean_content, target_style)
            previous_line_was_title = False

        # STANDARD TEXT / SECTION NAME CATCHER
        else:
            if previous_line_was_title:
                # If this line is right after "# SECTION...", it's the Name. Title it!
                add_safe_paragraph(stripped_text, STYLE_MAP["Title"])
            else:
                add_safe_paragraph(stripped_text, 'Normal')
            previous_line_was_title = False

    doc.save(output_path)
    print(f"✅ Success! Generated: {output_path}")

# ==========================================
# EXECUTION
# ==========================================
if __name__ == "__main__":
    # You can change these filenames here or load them from JSON if you prefer
    INPUT_MASTER = "office_master.docx"
    ARCHITECT_TEMPLATE = "architect_style.docx"
    OUTPUT_FILE = "Final_Project_Spec.docx"
    TEMP_MARKDOWN = "temp_converted.md"
    
    extract_master_to_markdown(INPUT_MASTER, TEMP_MARKDOWN)
    rebuild_from_markdown(TEMP_MARKDOWN, ARCHITECT_TEMPLATE, OUTPUT_FILE)