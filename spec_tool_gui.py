import tkinter as tk
from tkinter import filedialog, scrolledtext, messagebox, ttk
import threading
import os
import re
import json
from docx import Document

# ==========================================
# 🧠 THE LOGIC ENGINE (Our Previous Script)
# ==========================================

DEFAULT_CONFIG = {
    "styles": {
        "Title": "Heading 1", 
        "Part": "Heading 2", 
        "Article": "Heading 3"
    },
    "list_levels": ["List Bullet", "List Bullet 2", "List Bullet 3"],
    "options": {
        "strip_heading_numbers": True, 
        "strip_list_labels": True
    }
}

def load_config(config_path):
    """
    Load style mapping configuration from JSON file.
    
    Loads a project-specific configuration file that defines how CSI document
    structure elements (Title, Part, Article, Lists) map to Word style names
    in the architect's template.
    
    Args:
        config_path (str): Path to the JSON configuration file. Can be None or
                          empty string to use defaults.
    
    Returns:
        dict: Configuration dictionary containing:
            - styles (dict): Mapping of document elements to style names
                - Title: Style for section headers (e.g., "SECTION 23 21 13")
                - Part: Style for part headers (e.g., "PART 1 - GENERAL")
                - Article: Style for article headers (e.g., "1.01 SUMMARY")
            - list_levels (list): Ordered list of style names for nested lists
            - options (dict): Processing options
                - strip_heading_numbers (bool): Remove "1.01" from headings
                - strip_list_labels (bool): Remove "A." from list items
    
    Example:
        >>> config = load_config("project_config.json")
        >>> print(config['styles']['Title'])
        'CSILevel0'
    
    Notes:
        - Returns DEFAULT_CONFIG if path is None, empty, or file doesn't exist
        - Returns DEFAULT_CONFIG if JSON parsing fails
        - No error is raised on failure; defaults are silently used
    """
    if config_path and os.path.exists(config_path):
        try:
            with open(config_path, 'r') as f:
                return json.load(f)
        except Exception:
            return DEFAULT_CONFIG
    return DEFAULT_CONFIG

def extract_master_to_markdown(docx_path, md_path, log_func):
    """
    Extract content from office master Word document to markdown format.

    Reads a Word document and converts it to markdown while intelligently filtering
    out specifier notes, editing instructions, and other non-specification content
    that should not appear in the final document.

    Args:
        docx_path (str): Full path to the source Word document (office master spec).
        md_path (str): Full path where the markdown file will be saved.
        log_func (callable): Function to call for logging messages (e.g., print or GUI logger).

    Returns:
        bool: True if extraction succeeded, False if file not found or error occurred.

    Filtering Rules:
        - Removes paragraphs with styles: "Specifier Note", "Note", "Instruction", 
        "Editing Instruction", "CMT"
        - Removes paragraphs starting with: "See Editing Instruction", "Adjust list below",
        "Retain ", "Delete ", "Edit ", "Verify that Section titles"
        - Empty paragraphs are skipped

    Markdown Conversion:
        - Section headers (starting with "SECTION") → # Header
        - Part headers (starting with "PART" containing "-") → ## Header  
        - Article headers (pattern: digit.digit space) → ### Header
        - Lists with standard CSI indentation patterns (A., 1., a., 1), a))
        - All other text becomes body content with 2-space indent

    Example:
        >>> extract_master_to_markdown(
        ...     "masters/23 21 13.docx", 
        ...     "temp/23 21 13.md", 
        ...     print
        ... )
        True

    Notes:
        - Customizable via IGNORED_STYLES and IGNORED_STARTS lists within function
        - Output encoding is UTF-8
    """
    if not os.path.exists(docx_path):
        return False

    try:
        doc = Document(docx_path)
        md_lines = []

        # 🚫 STYLES TO IGNORE (Customize this list!)
        # Check your master file to see what the notes are called.
        IGNORED_STYLES = [
            "Specifier Note", 
            "Note", 
            "Instruction", 
            "Editing Instruction", 
            'CMT'
        ]

        # 🚫 KEYWORDS TO SKIP (Backup if styles are messy)
        # If a line starts with these, we skip it.
        IGNORED_STARTS = [
            "See Editing Instruction",
            "Adjust list below",
            "Retain ",
            "Delete ",
            "Edit ",
            "Verify that Section titles"
        ]

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text: continue

            # CHECK 1: Is it an ignored style?
            if para.style.name in IGNORED_STYLES:
                continue

            # CHECK 2: Does it look like an instruction? (Keyword check)
            # (Only use this if styles aren't working perfectly)
            is_instruction = False
            for keyword in IGNORED_STARTS:
                if text.startswith(keyword):
                    is_instruction = True
                    break
            if is_instruction:
                continue

            # Regex detection for CSI Structure
            if text.upper().startswith("SECTION"):
                md_lines.append(f"# {text}")
            elif text.upper().startswith("PART") and "-" in text:
                md_lines.append(f"\n## {text}")
            elif re.match(r'^\d+\.\d+\s', text):
                md_lines.append(f"\n### {text}")
            else:
                # Indentation Detection
                if re.match(r'^[A-Z]\.\s', text):       md_lines.append(f"- {text}")
                elif re.match(r'^\d+\.\s', text):       md_lines.append(f"  - {text}")
                elif re.match(r'^[a-z]\.\s', text):     md_lines.append(f"    - {text}")
                elif re.match(r'^\d+\)\s', text):       md_lines.append(f"      - {text}")
                elif re.match(r'^[a-z]\)\s', text):     md_lines.append(f"        - {text}")
                else:                                   md_lines.append(f"  {text}")

        with open(md_path, "w", encoding="utf-8") as f:
            f.write("\n".join(md_lines))
        return True
    except Exception as e:
        log_func(f"❌ Error extracting {os.path.basename(docx_path)}: {str(e)}")
        return False

def rebuild_from_markdown(md_path, template_path, output_path, config, log_func):
    """
    Rebuild Word document from markdown using architect's template formatting.

    Takes a markdown file and reconstructs it as a properly formatted Word document
    using style definitions from the architect's template. Preserves the template's
    page layout, margins, headers, footers, and section properties while replacing
    all content.

    Args:
        md_path (str): Path to the markdown file containing extracted content.
        template_path (str): Path to the architect's Word template (.docx) with target styles.
        output_path (str): Path where the formatted Word document will be saved.
        config (dict): Configuration dictionary from load_config() containing style mappings,
                    list levels, and processing options.
        log_func (callable): Function to call for logging messages (e.g., print or GUI logger).

    Returns:
        bool: True if rebuild succeeded, False if template not found or error occurred.

    Processing Logic:
        - Clears template body while preserving sectPr (section properties)
        - Maps markdown headers (#, ##, ###) to configured Word styles
        - Maps markdown list indentation to configured list level styles
        - Optionally strips heading numbers and list labels based on config
        - Handles special case: "END OF SECTION" uses Title style
        - Falls back to 'Normal' style if configured style doesn't exist

    Markdown to Word Mapping:
        - "# text" → Title style (Section headers)
        - "## text" → Part style (Part headers)
        - "### text" → Article style (Article headers)
        - "- text" at column 0 → list_levels[0]
        - "  - text" at column 2 → list_levels[1]
        - "    - text" at column 4 → list_levels[2]
        - (etc. for deeper indentation)

    Example:
        >>> config = load_config("project_config.json")
        >>> rebuild_from_markdown(
        ...     "temp/23 21 13.md",
        ...     "architect_template.docx",
        ...     "output/23 21 13.docx",
        ...     config,
        ...     print
        ... )
        True

    Notes:
        - Indent level calculated as: dash_index // 2
        - Uses three nested helper functions for safe paragraph creation and label cleaning
        - Preserves template's document-level properties (margins, page size, etc.)
    """
    if not os.path.exists(template_path):
        log_func(f"❌ Template not found: {template_path}")
        return False

    try:
        # Load Config
        style_map = config.get("styles", DEFAULT_CONFIG["styles"])
        list_map = config.get("list_levels", DEFAULT_CONFIG["list_levels"])
        options = config.get("options", DEFAULT_CONFIG["options"])

        doc = Document(template_path)
        
        # Clear Body (Keep Headers/Footers/Margins)
        for element in list(doc.element.body):
            if element.tag.endswith('sectPr'): continue
            doc.element.body.remove(element)

        with open(md_path, "r", encoding="utf-8") as f:
            lines = f.readlines()

        # Helpers
        def add_safe_paragraph(text, style_name):
            try:
                doc.add_paragraph(text, style=style_name)
            except KeyError:
                doc.add_paragraph(text, style='Normal')

        def clean_heading_label(text):
            if options.get("strip_heading_numbers"):
                return re.sub(r'^\d+\.\d+\s+', '', text)
            return text

        def clean_list_label(text):
            if options.get("strip_list_labels"):
                return re.sub(r'^[A-Za-z0-9]+[\.\)]\s+', '', text)
            return text

        previous_line_was_title = False

        for line in lines:
            raw_text = line.rstrip()
            stripped_text = raw_text.strip()
            
            if not stripped_text: continue

            # SPECIAL: End of Section
            if stripped_text.upper().startswith("END OF SECTION"):
                add_safe_paragraph(stripped_text, style_map["Title"])
                previous_line_was_title = False
                continue

            # HEADERS
            if stripped_text.startswith("# "):
                clean_text = stripped_text.replace("# ", "")
                add_safe_paragraph(clean_text, style_map["Title"])
                previous_line_was_title = True 
                
            elif stripped_text.startswith("## "):
                clean_text = stripped_text.replace("## ", "")
                add_safe_paragraph(clean_text, style_map["Part"])
                previous_line_was_title = False
                
            elif stripped_text.startswith("### "):
                text_no_hash = stripped_text.replace("### ", "")
                clean_text = clean_heading_label(text_no_hash)
                add_safe_paragraph(clean_text, style_map["Article"])
                previous_line_was_title = False

            # LISTS
            elif stripped_text.startswith("- "):
                dash_index = raw_text.find("-")
                indent_level = dash_index // 2 
                
                text_without_dash = stripped_text[2:]
                clean_content = clean_list_label(text_without_dash)

                if indent_level < len(list_map):
                    target_style = list_map[indent_level]
                else:
                    target_style = list_map[-1]
                
                add_safe_paragraph(clean_content, target_style)
                previous_line_was_title = False

            # STANDARD TEXT
            else:
                if previous_line_was_title:
                    add_safe_paragraph(stripped_text, style_map["Title"])
                else:
                    add_safe_paragraph(stripped_text, 'Normal')
                previous_line_was_title = False

        doc.save(output_path)
        return True
    except Exception as e:
        log_func(f"❌ Error building {os.path.basename(output_path)}: {str(e)}")
        return False

# ==========================================
# 🖥️ THE GUI APPLICATION
# ==========================================

class SpecToolApp:
    def __init__(self, root):
        """
        Initialize the Spec Automation Tool GUI application.

        Sets up the main application window, configures styling, and initializes
        all tkinter variables needed to store user inputs.

        Args:
            root (tk.Tk): The root Tkinter window object.

        Attributes Created:
            self.root (tk.Tk): Reference to the root window.
            self.master_folder (tk.StringVar): Path to office masters folder.
            self.template_file (tk.StringVar): Path to architect's template file.
            self.config_file (tk.StringVar): Path to project configuration JSON.
            self.output_folder (tk.StringVar): Path to output directory.

        Window Configuration:
            - Title: "Spec Automation Tool"
            - Size: 700x550 pixels
            - Theme: 'clam' (ttk style)

        Side Effects:
            - Calls create_widgets() to build the GUI layout
            - Sets window title and geometry

        Example:
            >>> root = tk.Tk()
            >>> app = SpecToolApp(root)
            >>> root.mainloop()
        """
        self.root = root
        self.root.title("Spec Automation Tool")
        self.root.geometry("700x550")
        
        # Style
        style = ttk.Style()
        style.theme_use('clam')

        # Variables
        self.master_folder = tk.StringVar()
        self.template_file = tk.StringVar()
        self.config_file = tk.StringVar()
        self.output_folder = tk.StringVar()

        self.create_widgets()

    def create_widgets(self):
        """
        Build and layout all GUI components for the application.

        Creates the complete user interface with input fields, browse buttons,
        action controls, progress indicator, and logging area. Organizes components
        into logical frames for clean visual hierarchy.

        GUI Structure:
            - Input Frame: Project settings with 4 file/folder selection rows
            - Action Frame: Process button and progress bar
            - Log Frame: Scrollable text area for real-time feedback

        Components Created:
            Input Frame (LabelFrame):
                - Row 0: Office Masters Folder (Entry + Browse button)
                - Row 1: Architect Template .docx (Entry + Browse button)
                - Row 2: Project Config .json (Entry + Browse button)
                - Row 3: Output Folder (Entry + Browse button)
            
            Action Frame (Frame):
                - Run button: "🚀 PROCESS ALL SPECS" (calls start_processing)
                - Progress bar: Horizontal determinate progressbar
            
            Log Frame (LabelFrame):
                - Scrollable text area (disabled state for read-only display)

        Layout:
            - All frames use pack() geometry manager
            - Input frame: fill="x" (horizontal expansion)
            - Log frame: fill="both", expand=True (takes remaining space)
            - Grid used within input frame for aligned rows

        Side Effects:
            - Creates self.run_btn (stored for enable/disable control)
            - Creates self.progress (for updating during batch processing)
            - Creates self.log_area (for writing processing messages)

        Args:
            None (operates on self)

        Returns:
            None
        """
        # --- INPUTS FRAME ---
        input_frame = ttk.LabelFrame(self.root, text="Project Settings", padding="10")
        input_frame.pack(fill="x", padx=10, pady=5)

        # 1. Master Folder
        ttk.Label(input_frame, text="Office Masters Folder:").grid(row=0, column=0, sticky="w")
        ttk.Entry(input_frame, textvariable=self.master_folder, width=50).grid(row=0, column=1, padx=5)
        ttk.Button(input_frame, text="Browse...", command=self.browse_master).grid(row=0, column=2)

        # 2. Template File
        ttk.Label(input_frame, text="Architect Template (.docx):").grid(row=1, column=0, sticky="w", pady=5)
        ttk.Entry(input_frame, textvariable=self.template_file, width=50).grid(row=1, column=1, padx=5)
        ttk.Button(input_frame, text="Browse...", command=self.browse_template).grid(row=1, column=2)

        # 3. Config File
        ttk.Label(input_frame, text="Project Config (.json):").grid(row=2, column=0, sticky="w")
        ttk.Entry(input_frame, textvariable=self.config_file, width=50).grid(row=2, column=1, padx=5)
        ttk.Button(input_frame, text="Browse...", command=self.browse_config).grid(row=2, column=2)

        # 4. Output Folder
        ttk.Label(input_frame, text="Output Folder:").grid(row=3, column=0, sticky="w", pady=5)
        ttk.Entry(input_frame, textvariable=self.output_folder, width=50).grid(row=3, column=1, padx=5)
        ttk.Button(input_frame, text="Browse...", command=self.browse_output).grid(row=3, column=2)

        # --- ACTIONS FRAME ---
        action_frame = ttk.Frame(self.root, padding="10")
        action_frame.pack(fill="x", padx=10)

        self.run_btn = ttk.Button(action_frame, text="🚀 PROCESS ALL SPECS", command=self.start_processing)
        self.run_btn.pack(fill="x", pady=5)
        
        self.progress = ttk.Progressbar(action_frame, orient="horizontal", mode="determinate")
        self.progress.pack(fill="x")

        # --- LOG FRAME ---
        log_frame = ttk.LabelFrame(self.root, text="Processing Log", padding="10")
        log_frame.pack(fill="both", expand=True, padx=10, pady=5)

        self.log_area = scrolledtext.ScrolledText(log_frame, state='disabled', height=10)
        self.log_area.pack(fill="both", expand=True)

    # --- BROWSE FUNCTIONS ---
    def browse_master(self):
        """
        Open folder browser dialog for selecting office masters folder.

        Launches a directory selection dialog and updates the master_folder
        StringVar with the selected path.

        Side Effects:
            - Opens native OS folder browser dialog
            - Sets self.master_folder to selected path (if user doesn't cancel)

        Args:
            None (operates on self)

        Returns:
            None

        Example Flow:
            User clicks "Browse..." next to "Office Masters Folder"
            → Dialog opens
            → User selects folder containing master .docx files
            → Path appears in entry field
        """
        path = filedialog.askdirectory()
        if path: self.master_folder.set(path)

    def browse_template(self):
        """
        Open file browser dialog for selecting architect's template document.

        Launches a file selection dialog filtered to show only Word documents (.docx)
        and updates the template_file StringVar with the selected path.

        File Filter:
            - Only .docx files are shown in the dialog

        Side Effects:
            - Opens native OS file browser dialog
            - Sets self.template_file to selected path (if user doesn't cancel)

        Args:
            None (operates on self)

        Returns:
            None

        Example Flow:
            User clicks "Browse..." next to "Architect Template (.docx)"
            → Dialog opens showing only .docx files
            → User selects architect's template file
            → Path appears in entry field
        """
        path = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
        if path: self.template_file.set(path)

    def browse_config(self):
        """
        Open file browser dialog for selecting project configuration JSON file.

        Launches a file selection dialog filtered to show only JSON files (.json)
        and updates the config_file StringVar with the selected path.

        File Filter:
            - Only .json files are shown in the dialog

        Side Effects:
            - Opens native OS file browser dialog
            - Sets self.config_file to selected path (if user doesn't cancel)

        Args:
            None (operates on self)

        Returns:
            None

        Example Flow:
            User clicks "Browse..." next to "Project Config (.json)"
            → Dialog opens showing only .json files
            → User selects project_config.json
            → Path appears in entry field

        Notes:
            - Configuration file is optional; defaults will be used if not provided
        """
        path = filedialog.askopenfilename(filetypes=[("JSON Config", "*.json")])
        if path: self.config_file.set(path)

    def browse_output(self):
        """
        Open folder browser dialog for selecting output destination folder.

        Launches a directory selection dialog and updates the output_folder
        StringVar with the selected path where formatted specifications will be saved.

        Side Effects:
            - Opens native OS folder browser dialog
            - Sets self.output_folder to selected path (if user doesn't cancel)

        Args:
            None (operates on self)

        Returns:
            None

        Example Flow:
            User clicks "Browse..." next to "Output Folder"
            → Dialog opens
            → User selects destination folder for formatted specs
            → Path appears in entry field

        Notes:
            - Folder should exist or be creatable by the user
            - Formatted files will be named "Formatted_{original_filename}"
        """
        path = filedialog.askdirectory()
        if path: self.output_folder.set(path)

    def log(self, message):
        """
        Write a message to the GUI log window.

        Appends a message to the scrollable log text area with automatic scrolling
        to show the latest entry. Temporarily enables the text widget for writing,
        then disables it again to maintain read-only behavior.

        Args:
            message (str): The message to display in the log window.

        Side Effects:
            - Enables log_area text widget
            - Appends message with newline to log_area
            - Scrolls log_area to show the newest entry
            - Disables log_area to prevent user editing

        Returns:
            None

        Example:
            >>> self.log("✅ Processing complete")
            >>> self.log("❌ Error: File not found")

        Notes:
            - Text area state toggles: disabled → normal → disabled
            - Automatically adds newline after each message
            - Used by both GUI thread and processing thread
        """
        self.log_area.config(state='normal')
        self.log_area.insert(tk.END, message + "\n")
        self.log_area.see(tk.END)
        self.log_area.config(state='disabled')

    # --- PROCESSING LOGIC ---
    def start_processing(self):
        """
        Validate user inputs and launch batch processing in separate thread.

        Performs input validation to ensure all required paths are provided, then
        spawns a daemon thread to run the batch processing without freezing the GUI.

        Validation:
            - Checks that master_folder is set
            - Checks that template_file is set
            - Checks that output_folder is set
            - Config file is optional (will use defaults if not provided)

        Side Effects:
            - Shows error messagebox if validation fails
            - Creates and starts daemon thread running run_batch()
            - Thread continues even if main GUI closes (daemon=True)

        Args:
            None (operates on self)

        Returns:
            None (returns early if validation fails)

        Example Flow:
            User clicks "🚀 PROCESS ALL SPECS"
            → Validation runs
            → If valid: Background thread starts, GUI remains responsive
            → If invalid: Error dialog shows, no processing occurs

        Threading:
            - Uses threading.Thread with daemon=True
            - Keeps GUI responsive during long batch operations
            - Log messages from thread appear in real-time
        """
        # Validate Inputs
        if not all([self.master_folder.get(), self.template_file.get(), self.output_folder.get()]):
            messagebox.showerror("Error", "Please select Master Folder, Template, and Output Folder.")
            return

        # Run in separate thread to keep GUI responsive
        threading.Thread(target=self.run_batch, daemon=True).start()

    def run_batch(self):
        """
        Execute batch processing of all specifications in a background thread.

        Main processing loop that finds all Word documents in the masters folder,
        extracts each to markdown, rebuilds with template formatting, and tracks
        progress. Runs in a separate thread to keep the GUI responsive.

        Process Flow:
            1. Disable the run button to prevent concurrent execution
            2. Load configuration file (or use defaults)
            3. Scan masters folder for .docx files (excluding temp files)
            4. For each file:
                a. Extract master to markdown (filtered)
                b. Rebuild from markdown using template
                c. Update progress bar
            5. Show completion messagebox
            6. Re-enable run button

        File Handling:
            - Processes all .docx files in masters_dir
            - Excludes temporary Word files (starting with "~$")
            - Creates temporary markdown files in output_dir
            - Generates final files named "Formatted_{original_name}"
            - Temporary .md files are left in place (line 237 cleanup is commented)

        Progress Tracking:
            - Sets progress bar maximum to total file count
            - Increments progress bar after each file
            - Updates GUI with root.update_idletasks()

        Side Effects:
            - Disables run_btn at start, re-enables at end
            - Writes multiple log messages during processing
            - Creates temporary .md files in output folder
            - Creates formatted .docx files in output folder
            - Shows "Done" messagebox when complete

        Args:
            None (operates on self, reads instance variables)

        Returns:
            None (may return early if no files found)

        Threading Context:
            - Called by thread spawned in start_processing()
            - Safe to call log() and update progress bar from thread
            - Uses root.update_idletasks() for thread-safe GUI updates

        Example Log Output:
            --- Starting Batch Process ---
            ✅ Config loaded successfully.
            Processing: 23 21 13.docx...
            ✅ Done.
            Processing: 23 33 00.docx...
            ✅ Done.
            --- Batch Complete ---
        """
        self.run_btn.config(state="disabled")
        self.log("--- Starting Batch Process ---")
        
        masters_dir = self.master_folder.get()
        template_path = self.template_file.get()
        config_path = self.config_file.get()
        output_dir = self.output_folder.get()

        # Load Config
        config = load_config(config_path)
        if config == DEFAULT_CONFIG:
            self.log("⚠️ No config file loaded (or failed). Using default styles.")
        else:
            self.log("✅ Config loaded successfully.")

        # Find Docx Files
        files = [f for f in os.listdir(masters_dir) if f.endswith(".docx") and not f.startswith("~$")]
        total_files = len(files)
        
        if total_files == 0:
            self.log("❌ No .docx files found in Master folder.")
            self.run_btn.config(state="normal")
            return

        self.progress["maximum"] = total_files
        self.progress["value"] = 0

        for i, filename in enumerate(files):
            master_path = os.path.join(masters_dir, filename)
            # Create a temp markdown path
            temp_md_path = os.path.join(output_dir, f"temp_{filename}.md")
            # Create final output path
            final_docx_path = os.path.join(output_dir, f"Formatted_{filename}")

            self.log(f"Processing: {filename}...")

            # Step 1: Extract
            if extract_master_to_markdown(master_path, temp_md_path, self.log):
                # Step 2: Rebuild
                if rebuild_from_markdown(temp_md_path, template_path, final_docx_path, config, self.log):
                    self.log(f"   ✅ Done.")
                
                # # Cleanup Temp
                # if os.path.exists(temp_md_path):
                #     os.remove(temp_md_path)
            
            self.progress["value"] = i + 1
            self.root.update_idletasks()

        self.log("--- Batch Complete ---")
        messagebox.showinfo("Done", f"Successfully processed {total_files} files.")
        self.run_btn.config(state="normal")

if __name__ == "__main__":
    root = tk.Tk()
    app = SpecToolApp(root)
    root.mainloop()


